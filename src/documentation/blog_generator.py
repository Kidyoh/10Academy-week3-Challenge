import json
import logging
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BlogPostGenerator:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.output_dir = Path('reports/blog')
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def generate_blog_post(self):
        """Generate complete blog post with all sections"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Optimizing Insurance Risk Analytics:', level=1)
        subtitle = doc.add_heading('A Deep Dive into AlphaCare\'s Data-Driven Transformation', level=2)
        
        # Add metadata
        self._add_metadata(doc)
        
        # Add content sections
        self._add_executive_summary(doc)
        self._add_introduction(doc)
        self._add_project_overview(doc)
        self._add_methodology(doc)
        self._add_findings(doc)
        self._add_technical_details(doc)
        self._add_recommendations(doc)
        self._add_conclusion(doc)
        self._add_footer(doc)
        
        # Save document
        output_file = self.output_dir / 'technical_blog_post.docx'
        doc.save(output_file)
        self.logger.info(f"Word document saved to {output_file}")

    def _add_metadata(self, doc):
        """Add metadata section to document"""
        date = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        tags = doc.add_paragraph("Tags: #DataScience #InsuranceAnalytics #MachineLearning #RiskAnalysis #Python")
        doc.add_paragraph()

    def _add_executive_summary(self, doc):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run("This technical case study explores how AlphaCare Insurance Solutions leveraged data analytics to optimize their risk assessment and marketing strategies. Through comprehensive analysis of historical insurance claim data from 2014-2015, we developed a sophisticated analytics pipeline that revealed significant insights into risk patterns and premium optimization opportunities.")
        doc.add_paragraph()

    def _add_introduction(self, doc):
        """Add introduction section"""
        doc.add_heading('1. Introduction & Business Context', level=1)
        
        # Role of Data Analytics
        doc.add_heading('The Role of Data Analytics in Insurance', level=2)
        p = doc.add_paragraph()
        p.add_run("The insurance industry is undergoing a dramatic transformation driven by data analytics. Traditional actuarial methods are being enhanced with machine learning and advanced statistical techniques, enabling more precise risk assessment and personalized premium pricing.")
        
        # Objectives
        doc.add_heading('AlphaCare\'s Objectives', level=2)
        objectives = doc.add_paragraph()
        objectives.add_run("• Develop cutting-edge risk and predictive analytics\n")
        objectives.add_run("• Optimize marketing strategies\n")
        objectives.add_run("• Identify low-risk customer segments\n")
        objectives.add_run("• Understand geographic risk patterns")

    def _add_project_overview(self, doc):
        """Add project overview section"""
        doc.add_heading('2. Project Overview', level=1)
        
        # Data Description
        doc.add_heading('Data Description', level=2)
        data_desc = doc.add_paragraph()
        data_desc.add_run("• Timeframe: February 2014 to August 2015\n")
        data_desc.add_run("• Source: Historical insurance claim data\n")
        data_desc.add_run("• Key components: Policy information, client demographics, vehicle details")

    def _add_methodology(self, doc):
        """Add methodology section"""
        doc.add_heading('3. Data Analysis & Methodology', level=1)
        
        # EDA
        doc.add_heading('Exploratory Data Analysis', level=2)
        self._add_code_snippet(doc, """
def analyze_data_quality(df):
    missing = pd.DataFrame({
        'Missing Values': df.isnull().sum(),
        'Percentage': (df.isnull().sum() / len(df) * 100).round(2)
    })
    return missing[missing['Missing Values'] > 0]
        """)
        
        # Statistical Modeling
        doc.add_heading('Statistical Modeling', level=2)
        self._add_code_snippet(doc, """
def fit_premium_models(df, feature_cols):
    X = df[feature_cols]
    y = df['TotalPremium']
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2)
    
    models = {
        'linear': LinearRegression(),
        'random_forest': RandomForestRegressor(),
        'xgboost': xgb.XGBRegressor()
    }
    return train_evaluate_models(models, X_train, X_test, y_train, y_test)
        """)

    def _add_findings(self, doc):
        """Add findings section"""
        doc.add_heading('4. Key Findings & Business Insights', level=1)
        
        # Risk Patterns
        doc.add_heading('Risk Patterns', level=2)
        patterns = doc.add_paragraph()
        patterns.add_run("• Significant variations in risk profiles across provinces\n")
        patterns.add_run("• Gender-based risk differences identified\n")
        patterns.add_run("• Vehicle type correlations with claim frequency")

    def _add_technical_details(self, doc):
        """Add technical details section"""
        doc.add_heading('5. Technical Implementation Details', level=1)
        
        # Pipeline Architecture
        doc.add_heading('Data Pipeline Architecture', level=2)
        self._add_code_snippet(doc, """
stages:
  prepare:
    cmd: python src/data/data_preprocessor.py
    deps:
      - data/raw/insurance_data.csv
    outs:
      - data/processed/preprocessed_data.csv
        """)

    def _add_recommendations(self, doc):
        """Add recommendations section"""
        doc.add_heading('6. Recommendations', level=1)
        
        # Business Strategy
        doc.add_heading('Business Strategy', level=2)
        strategy = doc.add_paragraph()
        strategy.add_run("1. Implement risk-based pricing strategy\n")
        strategy.add_run("2. Develop targeted marketing campaigns\n")
        strategy.add_run("3. Optimize customer acquisition costs")

    def _add_conclusion(self, doc):
        """Add conclusion section"""
        doc.add_heading('7. Conclusion', level=1)
        
        # Project Impact
        doc.add_heading('Project Impact', level=2)
        impact = doc.add_paragraph()
        impact.add_run("• Improved risk assessment accuracy\n")
        impact.add_run("• Data-driven decision making framework\n")
        impact.add_run("• Enhanced customer segmentation")

    def _add_footer(self, doc):
        """Add footer section"""
        doc.add_page_break()
        doc.add_heading('Author Information', level=1)
        author = doc.add_paragraph()
        author.add_run("Data Scientist at AlphaCare Insurance Solutions")

    def _add_code_snippet(self, doc, code):
        """Add formatted code snippet"""
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(code)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(9)

def main():
    logging.basicConfig(level=logging.INFO)
    generator = BlogPostGenerator()
    generator.generate_blog_post()

if __name__ == "__main__":
    main()